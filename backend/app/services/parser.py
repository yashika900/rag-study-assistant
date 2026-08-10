"""Document parsing service for PDF, TXT, and DOCX files."""

import traceback
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document
from PIL import Image
from pypdf import PdfReader

from backend.app.services.ocr import (
    OCRProcessingError,
    SUPPORTED_IMAGE_TYPES,
    extract_text_from_image,
)


class DocumentParsingError(Exception):
    """Raised when a file cannot be parsed into usable text."""


def parse_document(file_path: Path) -> list[Document]:
    """Extract text and basic metadata from a supported document."""

    print(f"Parser started for: {file_path}")
    extension = file_path.suffix.lower()
    try:
        if extension == ".pdf":
            documents = _parse_pdf(file_path)
        elif extension == ".txt":
            documents = _parse_txt(file_path)
        elif extension == ".docx":
            documents = _parse_docx(file_path)
        elif extension in SUPPORTED_IMAGE_TYPES:
            documents = extract_text_from_image(file_path)
        else:
            raise DocumentParsingError(f"Unsupported file type: {extension}")
    except DocumentParsingError:
        raise
    except Exception as exc:
        print("\n===== ERROR =====")
        print(str(exc))
        traceback.print_exc()
        raise DocumentParsingError(str(exc)) from exc

    if not any(doc.page_content.strip() for doc in documents):
        raise DocumentParsingError("The uploaded file does not contain readable text.")

    total_length = sum(len(document.page_content) for document in documents)
    print(f"Parser complete. Documents: {len(documents)}, text length: {total_length}")
    return documents


def _parse_pdf(file_path: Path) -> list[Document]:
    """Extract PDF text, falling back to OCR for scanned/image PDFs."""

    try:
        reader = PdfReader(str(file_path))
        print(f"PDF page count: {len(reader.pages)}")
        documents: list[Document] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(
                    Document(
                        page_content=_clean_text(text),
                        metadata={"source": file_path.name, "page": index},
                    )
                )

        if documents:
            print(f"PDF text extraction produced {len(documents)} readable pages")
            return documents

        print("No selectable PDF text found. Starting OCR fallback...")
        return _ocr_pdf(file_path)
    except Exception as exc:
        print("\n===== ERROR =====")
        print(str(exc))
        traceback.print_exc()
        raise DocumentParsingError("Could not read the PDF file.") from exc


def _ocr_pdf(file_path: Path) -> list[Document]:
    """Render scanned PDF pages to images and OCR them with Tesseract."""

    documents: list[Document] = []

    try:
        pdf = fitz.open(str(file_path))
        for page_index in range(len(pdf)):
            page_number = page_index + 1
            print(f"OCR processing PDF page {page_number}/{len(pdf)}")

            page = pdf.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)

            try:
                text = _ocr_image_object(image)
            except OCRProcessingError:
                continue

            if text.strip():
                documents.append(
                    Document(
                        page_content=_clean_text(text),
                        metadata={"source": file_path.name, "page": page_number},
                    )
                )
    except Exception as exc:
        print("\n===== ERROR =====")
        print(str(exc))
        traceback.print_exc()
        raise DocumentParsingError(f"PDF OCR failed: {exc}") from exc
    finally:
        try:
            pdf.close()
        except Exception:
            pass

    if not documents:
        raise DocumentParsingError(
            "Could not extract text from this PDF. It may be scanned, handwritten, or too low quality for OCR."
        )

    print(f"PDF OCR produced {len(documents)} readable pages")
    return documents


def _ocr_image_object(image: Image.Image) -> str:
    """OCR an in-memory image object."""

    try:
        import pytesseract

        text = pytesseract.image_to_string(image.convert("L")).strip()
    except Exception as exc:
        raise OCRProcessingError(f"OCR extraction failed: {exc}") from exc

    if not text:
        raise OCRProcessingError("No readable text found in image.")

    return text


def _parse_txt(file_path: Path) -> list[Document]:
    """Read a plain text file as a single document."""

    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")
    except Exception as exc:
        print("\n===== ERROR =====")
        print(str(exc))
        traceback.print_exc()
        raise DocumentParsingError("Could not read the text file.") from exc

    cleaned_text = _clean_text(text)
    if not cleaned_text:
        raise DocumentParsingError("Could not extract text from TXT")

    return [
        Document(
            page_content=cleaned_text,
            metadata={"source": file_path.name, "page": 0},
        )
    ]


def _parse_docx(file_path: Path) -> list[Document]:
    """Extract paragraph text from a DOCX file."""

    try:
        docx = DocxDocument(str(file_path))
        text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    except Exception as exc:
        print("\n===== ERROR =====")
        print(str(exc))
        traceback.print_exc()
        raise DocumentParsingError("Could not read the DOCX file.") from exc

    cleaned_text = _clean_text(text)
    if not cleaned_text:
        raise DocumentParsingError("Could not extract text from DOCX")

    return [
        Document(
            page_content=cleaned_text,
            metadata={"source": file_path.name, "page": 0},
        )
    ]


def _clean_text(text: str) -> str:
    """Normalize whitespace while preserving paragraph breaks."""

    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)
